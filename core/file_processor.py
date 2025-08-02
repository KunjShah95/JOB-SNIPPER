"""
Advanced File Processing Module for Job Snipper AI
Handles multiple file formats with security validation
"""
import os
import tempfile
import logging
from typing import Optional, Dict, Any
import PyPDF2
from docx import Document
import magic
import hashlib
from pathlib import Path

logger = logging.getLogger(__name__)

class FileProcessor:
    """Advanced file processor with security validation"""
    
    SUPPORTED_FORMATS = {
        'pdf': ['application/pdf'],
        'docx': ['application/vnd.openxmlformats-officedocument.wordprocessingml.document'],
        'doc': ['application/msword'],
        'txt': ['text/plain']
    }
    
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    def __init__(self):
        self.processed_files = {}
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from various file formats
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text content
        """
        try:
            # Validate file exists
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Check file size
            file_size = os.path.getsize(file_path)
            if file_size > self.MAX_FILE_SIZE:
                raise ValueError(f"File too large: {file_size} bytes (max: {self.MAX_FILE_SIZE})")
            
            # Detect file type
            file_extension = Path(file_path).suffix.lower().lstrip('.')
            
            # Extract text based on file type
            if file_extension == 'pdf':
                return self._extract_pdf_text(file_path)
            elif file_extension in ['doc', 'docx']:
                return self._extract_docx_text(file_path)
            elif file_extension == 'txt':
                return self._extract_txt_text(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise ValueError(f"Could not extract text from PDF: {e}")
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise ValueError(f"Could not extract text from DOCX: {e}")
    
    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
                
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read().strip()
            except Exception as e:
                logger.error(f"Error reading TXT file: {e}")
                raise ValueError(f"Could not read text file: {e}")
        except Exception as e:
            logger.error(f"Error extracting TXT text: {e}")
            raise ValueError(f"Could not extract text from TXT: {e}")
    
    def get_file_metadata(self, file_path: str) -> Dict[str, Any]:
        """Get comprehensive file metadata"""
        try:
            stat = os.stat(file_path)
            file_hash = self._calculate_file_hash(file_path)
            
            metadata = {
                'filename': os.path.basename(file_path),
                'size': stat.st_size,
                'created': stat.st_ctime,
                'modified': stat.st_mtime,
                'extension': Path(file_path).suffix.lower(),
                'hash': file_hash,
                'mime_type': self._get_mime_type(file_path)
            }
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error getting file metadata: {e}")
            return {}
    
    def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA256 hash of file"""
        try:
            hash_sha256 = hashlib.sha256()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_sha256.update(chunk)
            return hash_sha256.hexdigest()
        except Exception:
            return ""
    
    def _get_mime_type(self, file_path: str) -> str:
        """Get MIME type of file"""
        try:
            return magic.from_file(file_path, mime=True)
        except Exception:
            # Fallback to extension-based detection
            ext = Path(file_path).suffix.lower()
            mime_map = {
                '.pdf': 'application/pdf',
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                '.doc': 'application/msword',
                '.txt': 'text/plain'
            }
            return mime_map.get(ext, 'application/octet-stream')
    
    def validate_file_content(self, file_path: str) -> bool:
        """Validate that file contains meaningful content"""
        try:
            text = self.extract_text(file_path)
            
            # Check minimum length
            if len(text.strip()) < 50:
                return False
            
            # Check for common resume keywords
            resume_keywords = [
                'experience', 'education', 'skills', 'work', 'job',
                'university', 'college', 'degree', 'certification',
                'project', 'achievement', 'responsibility'
            ]
            
            text_lower = text.lower()
            keyword_count = sum(1 for keyword in resume_keywords if keyword in text_lower)
            
            # Should have at least 2 resume-related keywords
            return keyword_count >= 2
            
        except Exception:
            return False
    
    def clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        import re
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)]', ' ', text)
        
        # Remove extra spaces
        text = ' '.join(text.split())
        
        return text.strip()